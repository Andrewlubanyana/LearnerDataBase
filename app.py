import streamlit as st
import pandas as pd
import pdfplumber
from docx import Document
import plotly.express as px

st.set_page_config(page_title="Universal Student Database", layout="wide")

# --- 1. FILE HELPERS ---
def extract_pdf_table(file):
    with pdfplumber.open(file) as pdf:
        table = pdf.pages[0].extract_table()
        return pd.DataFrame(table[1:], columns=table[0]) if table else None

def extract_docx_table(file):
    doc = Document(file)
    if doc.tables:
        data = [[cell.text.strip() for cell in row.cells] for row in doc.tables[0].rows]
        return pd.DataFrame(data[1:], columns=data[0]) if data else None
    return None

# --- 2. APP UI ---
st.title("📊 Universal Student Database")
uploaded_file = st.file_uploader("Upload CSV, Excel, PDF, or Word", type=['csv', 'xlsx', 'pdf', 'docx'])

# Define df as None initially so the app doesn't crash
df = None

if uploaded_file:
    ext = uploaded_file.name.split('.')[-1].lower()
    try:
        if ext == 'csv':
            df = pd.read_csv(uploaded_file)
        elif ext == 'xlsx':
            df = pd.read_excel(uploaded_file)
        elif ext == 'pdf':
            df = extract_pdf_table(uploaded_file)
        elif ext == 'docx':
            df = extract_docx_table(uploaded_file)
        
        if df is not None:
            # CLEANING: Remove empty rows/cols and strip headers
            df = df.dropna(how='all')
            df.columns = [str(c).strip() for c in df.columns]
            
            # --- 3. SORTING & FILTERING (Only shows if df exists) ---
            st.sidebar.header("Settings")
            sort_by = st.sidebar.selectbox("Sort data by:", df.columns)
            order = st.sidebar.radio("Order:", ["Ascending", "Descending"])
            df = df.sort_values(by=sort_by, ascending=(order == "Ascending"))

            # --- 4. DISPLAY & ANALYZE ---
            st.subheader("Data Preview")
            st.dataframe(df, use_container_width=True)

            # Look for a "Mark" or "Score" column for math
            mark_col = next((c for c in df.columns if 'mark' in c.lower() or 'score' in c.lower()), None)
            if mark_col:
                df[mark_col] = pd.to_numeric(df[mark_col], errors='coerce')
                st.metric("Class Average", f"{df[mark_col].mean():.1f}%")
        else:
            st.warning("No table detected in this file. Please ensure data is in a table format.")

    except Exception as e:
        st.error(f"Error loading file: {e}")
else:
    st.info("👋 Welcome! Please upload a file to see the database and sorting options.")

        # --- 2. DATA CLEANING (Standardizing Column Names) ---
        # This makes the app "smart" by finding columns even if they are lowercase
        # --- ONLY RUN THIS IF DF WAS SUCCESSFULLY CREATED ---
        if df is not None:
            # This line removes hidden spaces from headers
            df.columns = [str(c).strip() for c in df.columns] 
            
            st.success(f"Loaded {uploaded_file.name} successfully!")
            
            # Now show the data
            st.dataframe(df, use_container_width=True)
            
            # ... rest of your analysis code ...
        else:
            st.error("The file was uploaded, but no data table could be extracted.") 
        
        # --- 3. SIDEBAR CONTROLS ---
        st.sidebar.header("Filter & Sort")
        
        # Sorting
        sort_by = st.sidebar.selectbox("Sort data by:", df.columns)
        order = st.sidebar.radio("Direction:", ["Ascending", "Descending"])
        
        df = df.sort_values(by=sort_by, ascending=(order == "Ascending"))

        # --- 4. ANALYSIS CALCULATIONS ---
        # We look for a column with 'mark' or 'score' in the name
        mark_col = next((c for c in df.columns if 'mark' in c.lower() or 'score' in c.lower()), None)

        if mark_col:
            avg_score = df[mark_col].mean()
            max_score = df[mark_col].max()
            min_score = df[mark_col].min()

            # Display "Metric" boxes at the top
            m1, m2, m3 = st.columns(3)
            m1.metric("Average Mark", f"{avg_score:.1f}%")
            m2.metric("Highest Mark", f"{max_score}%")
            m3.metric("Lowest Mark", f"{min_score}%")

        # --- 5. INTERACTIVE DATABASE TABLE ---
        st.subheader("Interactive Data Table")
        # Add a search bar for names
        search_term = st.text_input("Search by Name or Surname:")
        if search_term:
            # Filters the dataframe if the search term is in any column
            df = df[df.apply(lambda row: row.astype(str).str.contains(search_term, case=False).any(), axis=1)]
        
        st.dataframe(df, use_container_width=True)

        # --- 6. VISUAL CHARTS ---
        st.subheader("Data Visualizations")
        c1, c2 = st.columns(2)

        with c1:
            # Chart 1: Gender split (if column exists)
            gender_col = next((c for c in df.columns if 'gender' in c.lower()), None)
            if gender_col:
                fig1 = px.pie(df, names=gender_col, title="Gender Distribution", hole=0.4)
                st.plotly_chart(fig1)

        with c2:
            # Chart 2: Performance distribution
            if mark_col:
                fig2 = px.histogram(df, x=mark_col, title="Grade Spread", color_discrete_sequence=['#636EFA'])
                st.plotly_chart(fig2)

   # ... inside your file upload logic ...
    try:
        if file_extension == 'csv':
            df = pd.read_csv(uploaded_file)
        elif file_extension == 'xlsx':
            df = pd.read_excel(uploaded_file)
        # ... your other file types ...

        if df is not None:
            st.success("File Loaded!")
            st.dataframe(df)

    except Exception as e:
        st.error(f"Error processing file: {e}")
